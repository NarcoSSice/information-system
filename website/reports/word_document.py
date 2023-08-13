from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .models import *


def word_document():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    head = doc.add_paragraph()
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = head.add_run('Перелік наукових публікацій науково-педагогічних працівників кафедри (наукового підрозділу)')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph(style='List Number').add_run('Перелік монографій затверджених вченою радою університету').bold = True
    for pub in One.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number').add_run('Перелік колективних  монографій').bold = True
    for pub in Two.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number').add_run('Перелік підручників').bold = True
    for pub in Three.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number').add_run('Перелік навчальних посібників').bold = True
    for pub in Four.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number').add_run('Перелік наукових публікацій').bold = True
    doc.add_paragraph(style='List Number 2').add_run('Статті у фахових виданнях України ( зазначити категорію)').bold = True
    for pub in FiveOne.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number 2').add_run('Публікації у закордонних виданнях (окрім Російської федерації) із зазначенням виду публікації (стаття, тези)').bold = True
    doc.add_paragraph(style='List Number 3').add_run('Опубліковані:').bold = True
    for pub in FiveTwoOne.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number 3').add_run('Прийняті до друку:').bold = True
    for pub in FiveTwoTwo.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number 2').add_run('Наукові праці, опубліковані та підготовлені до друку у виданнях, що входять у МНБД (із зазначенням назви МНБД та виду публікації: стаття у журналі / збірнику наукових праць, матеріали конференцій тощо)').bold = True
    doc.add_paragraph(style='List Number 3').add_run('Опубліковані:').bold = True
    for pub in FiveThreeOne.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number 3').add_run('Прийняті до друку:').bold = True
    for pub in FiveThreeTwo.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number 2').add_run('Наукові праці, опубліковані та підготовлені до друку спільно із іноземними співавторами').bold = True
    doc.add_paragraph(style='List Number 3').add_run('Опубліковані:').bold = True
    for pub in FiveFourOne.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number 3').add_run('Прийняті до друку:').bold = True
    for pub in FiveFourTwo.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number').add_run('Публікації зі студентами ( зазначити стаття , тези)').bold = True
    for pub in Six.get_data():
        doc.add_paragraph().add_run(pub.name)
    doc.add_paragraph(style='List Number').add_run('Інші види (тези,  альбоми, атласи тощо)').bold = True
    for pub in Seven.get_data():
        doc.add_paragraph().add_run(pub.name)

    doc.save(r'reports\static\reports\word\Report.docx')
